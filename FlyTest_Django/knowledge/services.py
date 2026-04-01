"""
知识库服务模块
提供文档处理、向量化、检索等核心功能
"""

import os
import time
import hashlib
from typing import List, Dict, Any
import nltk
from django.conf import settings

# --- NLTK 数据路径配置 ---
# 将项目内部的 'nltk_data' 目录添加到 NLTK 的搜索路径中
# 这使得项目在任何环境中都能找到必要的数据，无需系统级安装
LOCAL_NLTK_DATA_PATH = os.path.join(settings.BASE_DIR, "nltk_data")
if os.path.exists(LOCAL_NLTK_DATA_PATH):
    if LOCAL_NLTK_DATA_PATH not in nltk.data.path:
        nltk.data.path.insert(0, LOCAL_NLTK_DATA_PATH)
        print(f"NLTK data path prepended with: {LOCAL_NLTK_DATA_PATH}")

# 设置完全离线模式，避免任何网络请求
os.environ["TRANSFORMERS_OFFLINE"] = "1"
os.environ["HF_DATASETS_OFFLINE"] = "1"
os.environ["HF_HUB_OFFLINE"] = "1"
os.environ["TOKENIZERS_PARALLELISM"] = "false"
# 禁用网络连接
os.environ["HF_HUB_DISABLE_TELEMETRY"] = "1"
os.environ["HF_HUB_DISABLE_PROGRESS_BARS"] = "1"
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"
# 设置极短的连接超时，强制快速失败
os.environ["HF_HUB_TIMEOUT"] = "1"
os.environ["REQUESTS_TIMEOUT"] = "1"
from django.conf import settings
from django.utils import timezone
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    UnstructuredPowerPointLoader,
    TextLoader,
    UnstructuredMarkdownLoader,
    UnstructuredHTMLLoader,
    WebBaseLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance,
    VectorParams,
    PointStruct,
    SparseVector,
    SparseVectorParams,
    SparseIndexParams,
    NamedVector,
    NamedSparseVector,
    models,
)
from langchain_core.documents import Document as LangChainDocument
from .models import (
    KnowledgeBase,
    Document,
    DocumentChunk,
    QueryLog,
    KnowledgeGlobalConfig,
)
import logging
import requests
import uuid
from typing import List, Optional, Dict
from langchain_core.embeddings import Embeddings

# 尝试导入 FastEmbed 用于 BM25 稀疏编码
# 注意：需要在导入前临时禁用离线模式
FASTEMBED_AVAILABLE = False
SparseTextEmbedding = None


def _init_fastembed():
    """延迟初始化 FastEmbed（避免模块级别的离线模式影响）"""
    global FASTEMBED_AVAILABLE, SparseTextEmbedding
    if FASTEMBED_AVAILABLE:
        return True

    # 临时禁用离线模式
    offline_vars = ["HF_HUB_OFFLINE", "TRANSFORMERS_OFFLINE", "HF_DATASETS_OFFLINE"]
    old_values = {var: os.environ.pop(var, None) for var in offline_vars}

    try:
        from fastembed import SparseTextEmbedding as _SparseTextEmbedding

        SparseTextEmbedding = _SparseTextEmbedding
        FASTEMBED_AVAILABLE = True
        return True
    except ImportError:
        return False
    finally:
        # 恢复环境变量
        for var, val in old_values.items():
            if val is not None:
                os.environ[var] = val


logger = logging.getLogger(__name__)


class SparseBM25Encoder:
    """基于 FastEmbed 的 BM25 稀疏编码器"""

    DEFAULT_MODEL = "Qdrant/bm25"

    def __init__(self, model_name: Optional[str] = None):
        # 初始化 FastEmbed（延迟导入）
        if not _init_fastembed():
            raise ImportError(
                "需要安装 fastembed 才能启用 BM25 稀疏向量: pip install fastembed"
            )

        self.model_name = model_name or self.DEFAULT_MODEL

        # 检查是否存在本地缓存（Docker 部署时模型已预下载）
        cache_path = os.environ.get(
            "FASTEMBED_CACHE_PATH", os.path.expanduser("~/.cache/fastembed")
        )
        model_cache_exists = (
            os.path.isdir(cache_path)
            and any("bm25" in d.lower() for d in os.listdir(cache_path))
            if os.path.exists(cache_path)
            else False
        )

        if model_cache_exists:
            # 有本地缓存时，保持离线模式，直接加载
            logger.info(f"📦 发现 BM25 模型缓存: {cache_path}，使用离线模式加载")
            self._encoder = SparseTextEmbedding(model_name=self.model_name)
            logger.info(f"✅ 初始化 BM25 稀疏编码器: {self.model_name}")
        else:
            # 无本地缓存时，临时禁用离线模式以下载模型
            offline_vars = [
                "HF_HUB_OFFLINE",
                "TRANSFORMERS_OFFLINE",
                "HF_DATASETS_OFFLINE",
            ]
            old_values = {var: os.environ.pop(var, None) for var in offline_vars}

            try:
                import huggingface_hub.constants


                if hasattr(huggingface_hub.constants, "HF_HUB_OFFLINE"):
                    huggingface_hub.constants.HF_HUB_OFFLINE = False
            except Exception:
                pass

            try:
                self._encoder = SparseTextEmbedding(model_name=self.model_name)
                logger.info(f"✅ 初始化 BM25 稀疏编码器: {self.model_name}")
            finally:
                for var, val in old_values.items():
                    if val is not None:
                        os.environ[var] = val

    def encode_documents(self, texts: List[str]) -> List:
        """编码文档列表"""
        return list(self._encoder.embed(texts))

    def encode_query(self, text: str):
        """编码查询"""
        results = list(self._encoder.query_embed(text))
        return results[0] if results else None


class CustomAPIEmbeddings(Embeddings):
    """自定义HTTP API嵌入服务"""



    def __init__(
        self,
        api_base_url: str,
        api_key: str = None,
        custom_headers: dict = None,
        model_name: str = "text-embedding",
    ):
        self.api_base_url = api_base_url.rstrip("/")
        self.api_key = api_key
        self.custom_headers = custom_headers or {}
        self.model_name = model_name

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """嵌入多个文档"""
        return [self.embed_query(text) for text in texts]

    def embed_query(self, text: str) -> List[float]:
        """嵌入单个查询"""
        headers = {"Content-Type": "application/json", **self.custom_headers}

        if self.api_key:
            headers["Authorization"] = f"Bearer {self.api_key}"

        data = {
            "input": text,
            "model": self.model_name,  # 使用配置的模型名
        }

        try:
            response = requests.post(
                self.api_base_url,  # 直接使用完整的API URL
                json=data,
                headers=headers,
                timeout=30,
            )
            response.raise_for_status()

            result = response.json()
            if "data" in result and len(result["data"]) > 0:
                return result["data"][0]["embedding"]
            else:
                raise ValueError(f"API响应格式错误: {result}")

        except Exception as e:
            raise RuntimeError(f"自定义API嵌入失败: {str(e)}")


class DocumentProcessor:
    """文档处理器 - 支持结构化解析"""

    def __init__(self):
        self.loaders = {
            "pdf": PyPDFLoader,
            "docx": self._load_docx_structured,  # 使用自定义结构化解析
            "doc": self._load_doc_structured,  # 支持旧版 .doc 格式
            "xlsx": self._load_excel_structured,  # Excel 表格
            "xls": self._load_excel_structured,  # 旧版 Excel
            "pptx": UnstructuredPowerPointLoader,
            "txt": TextLoader,
            "md": UnstructuredMarkdownLoader,
            "html": UnstructuredHTMLLoader,
        }

    def load_document(self, document: Document) -> List[LangChainDocument]:
        """加载文档内容"""
        try:
            logger.info(f"开始加载文档: {document.title} (ID: {document.id})")
            logger.info(f"文档类型: {document.document_type}")

            # 优先级：URL > 文本内容 > 文件
            if document.document_type == "url" and document.url:
                logger.info(f"从URL加载: {document.url}")
                return self._load_from_url(document.url)
            elif document.content:
                # 如果有文本内容，直接使用
                logger.info("从文本内容加载")
                return self._load_from_content(document.content, document.title)
            elif document.file and hasattr(document.file, "path"):
                file_path = document.file.path
                logger.info(f"从文件加载: {file_path}")

                # Windows路径兼容性处理
                if os.name == "nt":  # Windows系统
                    file_path = os.path.normpath(file_path)
                    if not os.path.isabs(file_path):
                        file_path = os.path.abspath(file_path)

                # 检查文件是否存在
                if os.path.exists(file_path):
                    logger.info(f"文件存在，开始加载: {file_path}")
                    return self._load_from_file(document)
                else:
                    raise FileNotFoundError(f"文件不存在: {file_path}")
            else:
                raise ValueError("文档没有可用的内容源（无URL、无文本内容、无文件）")
        except Exception as e:
            logger.error(f"加载文档失败 {document.id}: {e}")
            raise

    def _load_from_url(self, url: str) -> List[LangChainDocument]:
        """从URL加载文档"""
        loader = WebBaseLoader(url)
        return loader.load()

    def _load_from_content(self, content: str, title: str) -> List[LangChainDocument]:
        """从文本内容加载文档"""
        return [
            LangChainDocument(
                page_content=content, metadata={"source": title, "title": title}
            )
        ]

    def _load_from_file(self, document: Document) -> List[LangChainDocument]:
        """从文件加载文档"""
        file_path = document.file.path

        # Windows路径兼容性处理
        if os.name == "nt":  # Windows系统
            # 确保路径使用正确的分隔符
            file_path = os.path.normpath(file_path)
            # 转换为绝对路径
            if not os.path.isabs(file_path):
                file_path = os.path.abspath(file_path)

        logger.info(f"尝试加载文件: {file_path}")

        # 再次检查文件是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        loader = self.loaders.get(document.document_type)
        if not loader:
            raise ValueError(f"不支持的文档类型: {document.document_type}")

        try:
            # 检查是否为自定义方法（docx/doc 结构化解析）
            if callable(loader) and hasattr(loader, "__self__"):
                docs = loader(file_path, document)
            elif document.document_type == "txt":
                # 对于文本文件，使用UTF-8编码
                loader_instance = loader(file_path, encoding="utf-8")
                docs = loader_instance.load()
            else:
                # 其他类型使用标准 LangChain loader
                loader_instance = loader(file_path)
                docs = loader_instance.load()

            # 检查是否成功加载内容
            if not docs:
                raise ValueError(f"文档加载失败，没有内容: {file_path}")

            logger.info(f"成功加载文档，页数: {len(docs)}")

            # 添加元数据
            for doc in docs:
                doc.metadata.update(
                    {
                        "source": document.title,
                        "document_id": str(document.id),
                        "document_type": document.document_type,
                        "title": document.title,
                        "file_path": file_path,
                    }
                )

            return docs

        except Exception as e:
            logger.error(f"文档加载器失败: {e}")
            # 如果是文本文件，尝试直接读取
            if document.document_type == "txt":
                try:
                    logger.info("尝试直接读取文本文件...")
                    with open(file_path, "r", encoding="utf-8") as f:
                        content = f.read()

                    if not content.strip():
                        raise ValueError("文件内容为空")

                    return [
                        LangChainDocument(
                            page_content=content,
                            metadata={
                                "source": document.title,
                                "document_id": str(document.id),
                                "document_type": document.document_type,
                                "title": document.title,
                                "file_path": file_path,
                            },
                        )
                    ]
                except Exception as read_error:
                    logger.error(f"直接读取文件也失败: {read_error}")
                    raise
            else:
                raise

    def _load_docx_structured(
        self, file_path: str, document: Document
    ) -> List[LangChainDocument]:
        """结构化解析 .docx 文件，保留标题层级和表格结构"""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            logger.info(
                f"开始结构化解析 Word 文档，段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}"
            )

            # 创建元素到对象的映射
            paragraph_map = {p._element: p for p in doc.paragraphs}
            table_map = {t._element: t for t in doc.tables}

            content_parts = []
            extracted_paragraphs = 0
            extracted_tables = 0


            # 按文档顺序遍历所有元素
            for element in doc.element.body:
                if element.tag.endswith("p"):  # 段落
                    paragraph = paragraph_map.get(element)
                    if paragraph:
                        text = paragraph.text.strip()
                        if text:
                            markdown_text = self._convert_paragraph_to_markdown(
                                paragraph
                            )
                            content_parts.append(markdown_text)
                            extracted_paragraphs += 1

                elif element.tag.endswith("tbl"):  # 表格
                    table = table_map.get(element)
                    if table:
                        table_content = self._extract_table_content(table)
                        if table_content:
                            content_parts.append(table_content)
                            extracted_tables += 1

            content = "\n\n".join(content_parts)
            logger.info(
                f"Word 结构化解析完成 - 段落: {extracted_paragraphs}, 表格: {extracted_tables}, 内容长度: {len(content)}"
            )

            return [
                LangChainDocument(
                    page_content=content,
                    metadata={
                        "source": document.title,
                        "document_id": str(document.id),
                        "document_type": document.document_type,
                        "title": document.title,
                        "file_path": file_path,
                        "structured_parsing": True,
                        "paragraph_count": extracted_paragraphs,
                        "table_count": extracted_tables,
                    },
                )
            ]

        except Exception as e:
            logger.warning(f"结构化解析失败，降级为纯文本解析: {e}")
            # 降级为 Docx2txtLoader
            loader = Docx2txtLoader(file_path)
            docs = loader.load()
            for doc in docs:
                doc.metadata.update(
                    {
                        "source": document.title,
                        "document_id": str(document.id),
                        "document_type": document.document_type,
                        "title": document.title,
                        "file_path": file_path,
                        "structured_parsing": False,
                    }
                )
            return docs

    def _load_doc_structured(
        self, file_path: str, document: Document
    ) -> List[LangChainDocument]:
        """解析旧版 .doc 文件，优先转换为 docx 以保留结构"""
        import tempfile
        import subprocess

        # 检测文件真实格式
        with open(file_path, "rb") as f:
            header = f.read(8)

        # ZIP 魔数表示实际是 .docx
        if header[:4] == b"PK\x03\x04":
            logger.info("检测到 .doc 文件实际为 .docx 格式，使用 docx 解析器")
            return self._load_docx_structured(file_path, document)

        try:
            # 方法1: LibreOffice 转换为 docx
            with tempfile.TemporaryDirectory() as tmp_dir:
                result = subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        tmp_dir,
                        file_path,
                    ],
                    capture_output=True,
                    text=True,
                    timeout=120,
                )
                if result.returncode == 0:
                    import os



                    docx_file = os.path.join(
                        tmp_dir, os.path.basename(file_path).rsplit(".", 1)[0] + ".docx"
                    )
                    if os.path.exists(docx_file):
                        docs = self._load_docx_structured(docx_file, document)
                        logger.info(f"成功通过 LibreOffice 转换并解析 .doc 文件")
                        return docs
        except FileNotFoundError:
            logger.debug("LibreOffice 未安装")
        except Exception as e:
            logger.debug(f"LibreOffice 转换失败: {e}")

        # 方法2: antiword 提取纯文本并推断标题
        try:
            result = subprocess.run(
                ["antiword", "-w", "0", file_path],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0 and result.stdout.strip():
                content = self._infer_headings_from_plain_text(result.stdout.strip())
                logger.info(
                    f"成功通过 antiword 解析 .doc 文件，内容长度: {len(content)}"
                )
                return [
                    LangChainDocument(
                        page_content=content,
                        metadata={
                            "source": document.title,
                            "document_id": str(document.id),
                            "document_type": document.document_type,
                            "title": document.title,
                            "file_path": file_path,
                            "structured_parsing": False,
                            "parse_method": "antiword",
                        },
                    )
                ]
        except FileNotFoundError:
            logger.debug("antiword 未安装")
        except Exception as e:
            logger.debug(f"antiword 失败: {e}")

        raise ValueError(
            "无法解析 .doc 文件。请安装 LibreOffice 以获得最佳效果：\n"
            "Ubuntu/Debian: apt-get install libreoffice\n"
            "或者将文件另存为 .docx 格式后重新上传"
        )

    def _convert_paragraph_to_markdown(self, paragraph) -> str:
        """将 Word 段落转换为 Markdown 格式"""
        text = paragraph.text.strip()
        if not text:
            return ""

        style_name = paragraph.style.name if paragraph.style else ""

        # 根据样式转换为 Markdown 标题
        heading_map = {
            "heading 1": "# ",
            "heading 2": "## ",
            "heading 3": "### ",
            "heading 4": "#### ",
            "heading 5": "##### ",
            "heading 6": "###### ",
        }

        for style_key, prefix in heading_map.items():
            if style_key in style_name.lower():
                return f"{prefix}{text}"

        return text

    def _extract_table_content(self, table, depth=0) -> str:
        """提取表格内容为 Markdown 格式，支持合并单元格"""
        try:
            # 获取表格的实际行列数
            row_count = len(table.rows)
            if row_count == 0:
                return ""

            # 优先用 table.columns 获取列数（更稳定），失败时回退到行 cells 的最大长度
            try:
                max_cols = len(table.columns)
            except Exception:
                max_cols = max((len(row.cells) for row in table.rows), default=0)
            if max_cols == 0:
                return ""

            # 构建单元格网格，用于处理合并单元格
            # 网格元素格式：(单元格文本, 是否为合并延续单元格)
            grid = [[("", False) for _ in range(max_cols)] for _ in range(row_count)]

            def _sanitize_cell_text(text: str) -> str:
                """清理单元格文本，转义 Markdown 特殊字符"""
                text = (
                    (text or "")
                    .replace("\r", " ")
                    .replace("\n", " ")
                    .replace("\t", " ")
                )
                text = " ".join(text.split())  # 合并多个空格
                return text.replace("|", "\\|")  # 转义管道符

            for row_idx, row in enumerate(table.rows):
                # 记录当前行已处理的单元格（用于检测水平合并：同一行中重复引用同一 tc）
                processed_cells_in_row = set()

                for col_idx, cell in enumerate(row.cells):
                    if col_idx >= max_cols:
                        break

                    # 获取单元格的唯一标识
                    cell_id = id(cell._tc)

                    # 检测垂直合并
                    # Word 中 <w:vMerge/> 无 val 属性时通常表示"继续合并"
                    try:
                        tcPr = cell._tc.tcPr
                        v_merge = tcPr.vMerge if tcPr is not None else None
                        v_merge_val = (
                            getattr(v_merge, "val", None)
                            if v_merge is not None
                            else None
                        )

                        if v_merge is None:
                            is_v_merge_continue = False
                        elif v_merge_val == "restart":
                            is_v_merge_continue = False
                        elif v_merge_val == "continue":
                            is_v_merge_continue = True
                        else:
                            # val 为 None 时，通过检查上一行是否有 vMerge 来判断
                            if row_idx == 0:
                                is_v_merge_continue = False
                            else:
                                try:
                                    prev_cell = table.rows[row_idx - 1].cells[col_idx]
                                    prev_tcPr = prev_cell._tc.tcPr
                                    prev_v_merge = (
                                        prev_tcPr.vMerge
                                        if prev_tcPr is not None
                                        else None
                                    )
                                    is_v_merge_continue = prev_v_merge is not None
                                except Exception:
                                    is_v_merge_continue = False
                    except Exception:
                        is_v_merge_continue = False

                    # 垂直合并继续：标记为已合并
                    if is_v_merge_continue:
                        grid[row_idx][col_idx] = ("", True)
                        continue

                    # 水平合并继续：同一行中重复引用同一 tc
                    if cell_id in processed_cells_in_row:
                        grid[row_idx][col_idx] = ("", True)
                        continue

                    # 提取单元格内容
                    nested_tables = cell.tables
                    if nested_tables and depth < 3:
                        cell_text_parts = [
                            p.text.strip() for p in cell.paragraphs if p.text.strip()
                        ]
                        for nested_table in nested_tables:
                            nested_content = self._extract_table_content(
                                nested_table, depth + 1
                            )
                            if nested_content:
                                cell_text_parts.append(f"[嵌套表格] {nested_content}")
                        cell_text = _sanitize_cell_text(" ".join(cell_text_parts))
                    else:
                        cell_text = _sanitize_cell_text(cell.text.strip())

                    # 填充网格
                    grid[row_idx][col_idx] = (cell_text, False)
                    processed_cells_in_row.add(cell_id)

            # 生成 Markdown 表格
            table_rows = []
            for row_idx in range(row_count):
                row_cells = [grid[row_idx][col_idx][0] for col_idx in range(max_cols)]

                # 跳过全空行
                if not any(cell.strip() for cell in row_cells):
                    continue

                table_rows.append(" | ".join(row_cells))

                # 第一行后添加分隔符
                if len(table_rows) == 1:
                    separator = " | ".join(["---"] * max_cols)
                    table_rows.append(separator)

            if table_rows:
                return "\n".join(table_rows)
            return ""

        except Exception as e:
            logger.warning(f"表格提取失败: {e}")
            return ""

    def _infer_headings_from_plain_text(self, content: str) -> str:
        """从纯文本推断标题结构"""
        import re

        lines = content.split("\n")
        result_lines = []

        # 常见标题模式
        patterns = [
            (r"^第[一二三四五六七八九十百]+[章节部分]\s*", 1),
            (r"^[一二三四五六七八九十]+[、.．]\s*", 2),
            (r"^[（\(][一二三四五六七八九十]+[）\)]\s*", 3),
            (r"^(\d+)\s*[、.．]\s*", 2),
            (r"^(\d+\.\d+)\s+", 3),
            (r"^(\d+\.\d+\.\d+)\s+", 4),
        ]

        for line in lines:
            stripped = line.strip()
            if not stripped:
                result_lines.append("")
                continue

            matched = False
            for pattern, level in patterns:
                if re.match(pattern, stripped) and len(stripped) <= 80:
                    prefix = "#" * level + " "
                    result_lines.append(prefix + stripped)
                    matched = True
                    break

            if not matched:
                result_lines.append(line)

        return "\n".join(result_lines)

    def _load_excel_structured(
        self, file_path: str, document: Document
    ) -> List[LangChainDocument]:
        """解析 Excel 文件（.xlsx/.xls），将每个工作表转换为 Markdown 表格"""
        try:
            import pandas as pd

            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names

            logger.info(f"开始解析 Excel 文件，工作表数量: {len(sheet_names)}")

            content_parts = []
            total_rows = 0

            for sheet_name in sheet_names:
                try:
                    # 读取工作表，保留所有数据
                    df = pd.read_excel(excel_file, sheet_name=sheet_name, dtype=str)
                    df = df.fillna("")  # 空值替换为空字符串

                    if df.empty:
                        continue

                    total_rows += len(df)

                    # 生成工作表标题
                    content_parts.append(f"## {sheet_name}")

                    # 转换为 Markdown 表格
                    markdown_table = self._dataframe_to_markdown(df)
                    if markdown_table:
                        content_parts.append(markdown_table)

                except Exception as e:
                    logger.warning(f"解析工作表 '{sheet_name}' 失败: {e}")
                    continue

            content = "\n\n".join(content_parts)
            logger.info(
                f"Excel 解析完成 - 工作表: {len(sheet_names)}, 总行数: {total_rows}, 内容长度: {len(content)}"
            )

            return [
                LangChainDocument(
                    page_content=content,
                    metadata={
                        "source": document.title,
                        "document_id": str(document.id),
                        "document_type": document.document_type,
                        "title": document.title,
                        "file_path": file_path,
                        "structured_parsing": True,
                        "sheet_count": len(sheet_names),
                        "total_rows": total_rows,
                    },
                )
            ]

        except ImportError:
            raise ValueError(
                "需要安装 pandas 和 openpyxl: pip install pandas openpyxl xlrd"
            )
        except Exception as e:
            logger.error(f"Excel 解析失败: {e}")
            raise ValueError(f"无法解析 Excel 文件: {e}")

    def _dataframe_to_markdown(self, df) -> str:
        """将 DataFrame 转换为 Markdown 表格"""
        if df.empty:
            return ""

        def _sanitize(text):
            """清理单元格文本"""
            text = str(text).replace("\r", " ").replace("\n", " ").replace("\t", " ")
            text = " ".join(text.split())
            return text.replace("|", "\\|")

        # 表头
        headers = [_sanitize(col) for col in df.columns]
        header_row = " | ".join(headers)
        separator = " | ".join(["---"] * len(headers))

        # 数据行
        data_rows = []
        for _, row in df.iterrows():
            cells = [_sanitize(cell) for cell in row]
            data_rows.append(" | ".join(cells))

        # 组合表格
        table_parts = [header_row, separator] + data_rows
        return "\n".join(table_parts)


class VectorStoreManager:
    """向量存储管理器 - 支持稠密+稀疏混合检索 + Reranker精排"""

    # 向量名称常量
    DENSE_VECTOR_NAME = "dense"
    SPARSE_VECTOR_NAME = "bm25"
    # RRF 融合参数
    RRF_K = 60
    # Reranker 配置
    RERANKER_MODEL = "bge-reranker-v2-m3"
    RERANKER_ENABLED = True  # 可通过环境变量控制

    # 类级别的缓存
    _vector_store_cache = {}
    _embeddings_cache = {}
    _sparse_encoder_cache = {}
    _reranker_config_cache = None
    _reranker_config_cache_time = 0
    _global_config_cache = None
    _global_config_cache_time = 0

    def __init__(self, knowledge_base: KnowledgeBase):
        self.knowledge_base = knowledge_base
        self.global_config = self._get_global_config()
        self.embeddings = self._get_embeddings_instance()
        self.sparse_encoder = self._get_sparse_encoder()
        self._log_embedding_info()

    @classmethod
    def _get_global_config(cls):
        """获取全局配置（带缓存，5分钟过期）"""
        import time

        current_time = time.time()

        # 缓存5分钟
        if (
            cls._global_config_cache
            and (current_time - cls._global_config_cache_time) < 300
        ):
            return cls._global_config_cache

        cls._global_config_cache = KnowledgeGlobalConfig.get_config()
        cls._global_config_cache_time = current_time
        return cls._global_config_cache

    @classmethod
    def clear_global_config_cache(cls):
        """清理全局配置缓存"""
        cls._global_config_cache = None
        cls._global_config_cache_time = 0

    def _get_embeddings_instance(self):
        """获取嵌入模型实例，使用全局配置"""
        config = self.global_config
        cache_key = (
            f"{config.embedding_service}_{config.api_base_url}_{config.model_name}"
        )

        if cache_key not in self._embeddings_cache:
            embedding_service = config.embedding_service

            try:
                if embedding_service == "openai":
                    self._embeddings_cache[cache_key] = self._create_openai_embeddings(
                        config
                    )
                elif embedding_service == "azure_openai":
                    self._embeddings_cache[cache_key] = self._create_azure_embeddings(
                        config
                    )
                elif embedding_service == "ollama":
                    self._embeddings_cache[cache_key] = self._create_ollama_embeddings(
                        config
                    )
                elif embedding_service == "xinference":
                    self._embeddings_cache[cache_key] = (
                        self._create_xinference_embeddings(config)
                    )
                elif embedding_service == "custom":
                    self._embeddings_cache[cache_key] = (
                        self._create_custom_api_embeddings(config)
                    )
                else:
                    raise ValueError(f"不支持的嵌入服务: {embedding_service}")

                # 测试嵌入功能
                test_embedding = self._embeddings_cache[cache_key].embed_query(
                    "模型功能测试"
                )
                logger.info(
                    f"✅ 嵌入模型测试成功: {embedding_service}, 维度: {len(test_embedding)}"
                )

            except Exception as e:
                logger.error(f"❌ 嵌入服务 {embedding_service} 初始化失败: {str(e)}")
                raise

        return self._embeddings_cache[cache_key]

    def _get_sparse_encoder(self) -> Optional[SparseBM25Encoder]:
        """获取 BM25 稀疏编码器（带缓存）"""
        cache_key = self.SPARSE_VECTOR_NAME

        if cache_key not in self._sparse_encoder_cache:
            try:
                self._sparse_encoder_cache[cache_key] = SparseBM25Encoder()
            except ImportError as e:
                logger.warning(f"⚠️ FastEmbed 未安装，将使用纯稠密向量检索: {e}")
                self._sparse_encoder_cache[cache_key] = None
            except Exception as e:
                logger.warning(f"⚠️ BM25 编码器初始化失败: {e}，降级为纯稠密检索")
                self._sparse_encoder_cache[cache_key] = None

        return self._sparse_encoder_cache[cache_key]

    def _create_openai_embeddings(self, config):
        """创建OpenAI Embeddings实例"""
        try:
            from langchain_openai import OpenAIEmbeddings
        except ImportError:
            raise ImportError("需要安装langchain-openai: pip install langchain-openai")

        kwargs = {
            "model": config.model_name or "text-embedding-ada-002",
        }

        if config.api_key:
            kwargs["api_key"] = config.api_key
        if config.api_base_url:
            kwargs["base_url"] = config.api_base_url

        logger.info(f"🚀 初始化OpenAI嵌入模型: {kwargs['model']}")
        return OpenAIEmbeddings(**kwargs)

    def _create_azure_embeddings(self, config):
        """创建Azure OpenAI Embeddings实例"""
        try:
            from langchain_openai import AzureOpenAIEmbeddings
        except ImportError:
            raise ImportError("需要安装langchain-openai: pip install langchain-openai")

        if not all([config.api_key, config.api_base_url]):
            raise ValueError("Azure OpenAI需要配置api_key和api_base_url")

        kwargs = {
            "model": config.model_name or "text-embedding-ada-002",
            "api_key": config.api_key,
            "azure_endpoint": config.api_base_url,
            "api_version": "2024-02-15-preview",
        }


        kwargs["deployment"] = config.model_name or "text-embedding-ada-002"

        logger.info(f"🚀 初始化Azure OpenAI嵌入模型: {kwargs['model']}")
        return AzureOpenAIEmbeddings(**kwargs)

    def _create_ollama_embeddings(self, config):
        """创建Ollama Embeddings实例"""
        try:
            from langchain_ollama import OllamaEmbeddings
        except ImportError:
            raise ImportError("需要安装langchain-ollama: pip install langchain-ollama")

        kwargs = {
            "model": config.model_name or "bge-m3",
        }

        if config.api_base_url:
            kwargs["base_url"] = config.api_base_url
        else:
            kwargs["base_url"] = "http://localhost:11434"

        logger.info(f"🚀 初始化Ollama嵌入模型: {kwargs['model']}")
        return OllamaEmbeddings(**kwargs)

    def _create_xinference_embeddings(self, config):
        """创建Xinference Embeddings实例"""
        if not config.api_base_url:
            base_url = "http://localhost:9997"
        else:
            base_url = config.api_base_url.rstrip("/")

        logger.info(f"🚀 初始化Xinference嵌入模型: {config.model_name or 'bge-m3'}")
        return CustomAPIEmbeddings(
            api_base_url=f"{base_url}/v1/embeddings",
            api_key=config.api_key or "",
            custom_headers={},
            model_name=config.model_name or "bge-m3",
        )

    def _get_reranker_config(self) -> tuple:
        """获取 Reranker 配置（独立于 Embedding）"""
        config = self.global_config

        # 检查是否启用 Reranker
        reranker_service = getattr(config, "reranker_service", "none")
        if reranker_service == "none":
            return None, None

        # 获取 Reranker API 地址
        reranker_api_url = getattr(config, "reranker_api_url", None)
        if not reranker_api_url:
            # 如果未配置独立地址，尝试使用 Embedding 服务地址（仅限 Xinference）
            if config.embedding_service == "xinference" and config.api_base_url:
                reranker_api_url = config.api_base_url
            elif reranker_service == "xinference":
                reranker_api_url = "http://localhost:9997"
            else:
                return None, None

        # 获取模型名称
        reranker_model = getattr(config, "reranker_model_name", "bge-reranker-v2-m3")

        base_url = reranker_api_url.rstrip("/")
        return f"{base_url}/v1/rerank", reranker_model

    def _get_reranker_url(self) -> Optional[str]:
        """获取 Reranker 服务地址（带缓存，跟随全局配置缓存过期）"""
        url, _ = self._get_reranker_config()
        return url

    def _get_reranker_model(self) -> str:
        """获取 Reranker 模型名称"""
        _, model = self._get_reranker_config()
        return model or self.RERANKER_MODEL

    def _rerank(
        self, query: str, candidates: List[Dict[str, Any]], top_k: int
    ) -> List[Dict[str, Any]]:
        """使用 Reranker 对候选结果进行精排"""
        reranker_url = self._get_reranker_url()
        reranker_model = self._get_reranker_model()
        if not reranker_url or not candidates:
            return candidates[:top_k]

        try:
            import requests as http_requests

            # 准备文档列表
            documents = [
                c.get("payload", {}).get("page_content", "") for c in candidates
            ]
            if not any(documents):
                return candidates[:top_k]

            # 调用 Reranker API
            logger.info(
                f"🔄 Reranker 请求: URL={reranker_url}, model={reranker_model}, docs={len(documents)}"
            )
            response = http_requests.post(
                reranker_url,
                json={
                    "model": reranker_model,
                    "query": query,
                    "documents": documents,
                    "top_n": top_k,
                },
                timeout=30,
            )

            if not response.ok:
                logger.warning(
                    f"⚠️ Reranker 调用失败: HTTP {response.status_code} - {response.text[:200]}, 降级为 RRF 排序"
                )
                return candidates[:top_k]

            rerank_result = response.json()
            results = rerank_result.get("results", [])
            logger.info(
                f"🔄 Reranker 原始返回: {len(results)} 条, 分数范围: {[r.get('relevance_score', 0) for r in results[:3]]}"
            )

            if not results:
                logger.warning("⚠️ Reranker 返回空结果，降级为 RRF 排序")
                return candidates[:top_k]

            # 根据 rerank 结果重新排序
            reranked = []
            for item in results:
                idx = item.get("index", 0)
                rerank_score = item.get("relevance_score", 0.0)
                if 0 <= idx < len(candidates):
                    candidate = candidates[idx].copy()
                    candidate["rerank_score"] = rerank_score
                    reranked.append(candidate)

            logger.info(f"🎯 Reranker 精排完成: {len(reranked)} 条结果")
            return reranked

        except Exception as e:
            logger.warning(f"⚠️ Reranker 调用异常: {e}, 降级为 RRF 排序")
            return candidates[:top_k]

    def _create_custom_api_embeddings(self, config):
        """创建自定义API Embeddings实例"""
        if not config.api_base_url:
            raise ValueError("自定义API需要配置api_base_url")

        logger.info(f"🚀 初始化自定义API嵌入模型: {config.api_base_url}")
        return CustomAPIEmbeddings(
            api_base_url=config.api_base_url,
            api_key=config.api_key,
            custom_headers={},
            model_name=config.model_name,
        )

    def _log_embedding_info(self):
        """记录嵌入模型信息"""
        embedding_type = type(self.embeddings).__name__
        config = self.global_config
        logger.info(f"   🌟 知识库: {self.knowledge_base.name}")
        logger.info(f"   🎯 配置的嵌入模型: {config.model_name}")
        logger.info(f"   ✅ 实际使用的嵌入模型: {embedding_type}")

        if embedding_type == "OpenAIEmbeddings":
            logger.info(f"   🎉 说明: 使用OpenAI嵌入API服务")
        elif embedding_type == "AzureOpenAIEmbeddings":
            logger.info(f"   🎉 说明: 使用Azure OpenAI嵌入API服务")
        elif embedding_type == "OllamaEmbeddings":
            logger.info(f"   🎉 说明: 使用Ollama本地API嵌入服务")
        elif embedding_type == "CustomAPIEmbeddings":
            if config.embedding_service == "xinference":
                logger.info(f"   🎉 说明: 使用Xinference嵌入服务（支持Reranker）")
            else:
                logger.info(f"   🎉 说明: 使用自定义HTTP API嵌入服务")

        self._vector_store = None
        self._qdrant_client = None
        logger.info(f"🤖 向量存储管理器初始化完成:")
        logger.info(
            f"   📋 知识库: {self.knowledge_base.name} (ID: {self.knowledge_base.id})"
        )
        logger.info(f"   🎯 配置的嵌入模型: {config.model_name}")
        logger.info(f"   ✅ 实际使用的嵌入模型: {embedding_type}")
        logger.info(f"   💾 向量存储类型: Qdrant")

    def _get_qdrant_url(self) -> str:
        """获取 Qdrant 服务地址"""
        return os.environ.get("QDRANT_URL", "http://localhost:8918")

    def _get_collection_name(self) -> str:
        """获取集合名称"""
        return f"kb_{self.knowledge_base.id}"

    @property
    def qdrant_client(self) -> QdrantClient:
        """获取 Qdrant 客户端"""
        if self._qdrant_client is None:
            qdrant_url = self._get_qdrant_url()
            self._qdrant_client = QdrantClient(url=qdrant_url)
            logger.info(f"🔗 已连接 Qdrant: {qdrant_url}")
        return self._qdrant_client

    @property
    def vector_store(self):
        """获取向量存储实例（带缓存和健康检查）"""
        if self._vector_store is None:
            cache_key = str(self.knowledge_base.id)

            if cache_key in self._vector_store_cache:
                cached_store = self._vector_store_cache[cache_key]
                try:
                    # 验证 Qdrant 集合是否存在
                    self.qdrant_client.get_collection(self._get_collection_name())
                    logger.info(f"使用缓存的向量存储实例: {cache_key}")
                    self._vector_store = cached_store
                except Exception as e:
                    logger.warning(f"缓存的 Collection 无效,重新创建: {e}")
                    del self._vector_store_cache[cache_key]
                    logger.info(f"创建新的向量存储实例: {cache_key}")
                    self._vector_store = self._create_vector_store()
                    self._vector_store_cache[cache_key] = self._vector_store
            else:
                logger.info(f"创建新的向量存储实例: {cache_key}")
                self._vector_store = self._create_vector_store()
                self._vector_store_cache[cache_key] = self._vector_store

        return self._vector_store

    @classmethod
    def clear_cache(cls, knowledge_base_id=None):
        """清理向量存储缓存"""
        if knowledge_base_id:
            cache_key = str(knowledge_base_id)
            if cache_key in cls._vector_store_cache:
                del cls._vector_store_cache[cache_key]
                logger.info(f"已清理知识库 {cache_key} 的向量存储缓存")

            # 清理 Qdrant 集合
            try:
                qdrant_url = os.environ.get("QDRANT_URL", "http://localhost:8918")
                client = QdrantClient(url=qdrant_url)
                collection_name = f"kb_{knowledge_base_id}"
                if client.collection_exists(collection_name):
                    client.delete_collection(collection_name)
                    logger.info(f"已删除 Qdrant 集合: {collection_name}")
            except Exception as e:
                logger.warning(f"清理 Qdrant 集合失败: {e}")
        else:
            # 清理所有缓存
            cls._vector_store_cache.clear()
            cls._embeddings_cache.clear()
            cls._sparse_encoder_cache.clear()
            logger.info("已清理所有向量存储缓存")

    def _create_vector_store(self):
        """创建 Qdrant 向量存储（支持稠密+稀疏混合）"""
        collection_name = self._get_collection_name()

        # 获取嵌入向量维度
        test_embedding = self.embeddings.embed_query("测试")
        vector_size = len(test_embedding)

        # 配置命名向量（用于混合检索）
        vectors_config = {
            self.DENSE_VECTOR_NAME: VectorParams(
                size=vector_size, distance=Distance.COSINE
            )
        }

        # 配置稀疏向量
        sparse_vectors_config = None
        if self.sparse_encoder:
            sparse_vectors_config = {
                self.SPARSE_VECTOR_NAME: SparseVectorParams(
                    index=SparseIndexParams(on_disk=False)
                )
            }

        # 确保集合存在
        try:
            if not self.qdrant_client.collection_exists(collection_name):
                self.qdrant_client.create_collection(
                    collection_name=collection_name,
                    vectors_config=vectors_config,
                    sparse_vectors_config=sparse_vectors_config,
                )
                mode = "稀疏+稠密混合" if sparse_vectors_config else "纯稠密"
                logger.info(
                    f"✅ 创建 Qdrant 集合: {collection_name}, 维度: {vector_size}, 模式: {mode}"
                )
            else:
                # 检查是否需要更新稀疏配置
                if sparse_vectors_config:
                    try:
                        self.qdrant_client.update_collection(
                            collection_name=collection_name,
                            sparse_vectors_config=sparse_vectors_config,
                        )
                    except Exception as e:
                        logger.debug(f"跳过稀疏配置更新: {e}")
        except Exception as e:
            logger.warning(f"检查/创建集合时出错: {e}")

        # 使用 LangChain 的 QdrantVectorStore（用于兼容性，实际混合查询直接用 client）
        qdrant_store = QdrantVectorStore(
            client=self.qdrant_client,
            collection_name=collection_name,
            embedding=self.embeddings,
            vector_name=self.DENSE_VECTOR_NAME,
        )

        return qdrant_store

    def add_documents(
        self, documents: List[LangChainDocument], document_obj: Document
    ) -> List[str]:
        """添加文档到向量存储（稠密+稀疏混合）"""
        try:
            # 确保集合存在（触发 vector_store 属性会创建集合）
            _ = self.vector_store

            # 文档分块
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.knowledge_base.chunk_size,
                chunk_overlap=self.knowledge_base.chunk_overlap,
            )
            chunks = text_splitter.split_documents(documents)

            # 生成唯一的 vector_ids
            vector_ids = [str(uuid.uuid4()) for _ in chunks]
            chunk_texts = [chunk.page_content for chunk in chunks]

            # 计算稠密向量
            dense_embeddings = self.embeddings.embed_documents(chunk_texts)

            # 计算稀疏向量（如果可用）
            sparse_embeddings = None
            if self.sparse_encoder:
                sparse_embeddings = self.sparse_encoder.encode_documents(chunk_texts)

            # 构建 PointStruct 列表
            points: List[PointStruct] = []
            for i, (chunk, vector_id, dense_vector) in enumerate(
                zip(chunks, vector_ids, dense_embeddings)
            ):
                payload = dict(chunk.metadata or {})
                payload.update(
                    {
                        "page_content": chunk.page_content,
                        "document_id": str(document_obj.id),
                        "chunk_index": i,
                        "vector_id": vector_id,
                        "knowledge_base_id": str(self.knowledge_base.id),
                    }
                )

                # 构建向量配置
                vectors = {self.DENSE_VECTOR_NAME: dense_vector}

                # 添加稀疏向量（如果可用）
                sparse_vectors = None
                if sparse_embeddings and sparse_embeddings[i]:
                    sparse_vec = sparse_embeddings[i]
                    sparse_vectors = {
                        self.SPARSE_VECTOR_NAME: SparseVector(
                            indices=sparse_vec.indices.tolist(),
                            values=sparse_vec.values.tolist(),
                        )
                    }

                point = PointStruct(
                    id=vector_id,
                    vector=vectors,
                    payload=payload,
                )

                # Qdrant SDK 需要单独设置 sparse_vectors
                if sparse_vectors:
                    point = PointStruct(
                        id=vector_id,
                        vector={
                            self.DENSE_VECTOR_NAME: dense_vector,
                            self.SPARSE_VECTOR_NAME: SparseVector(
                                indices=sparse_embeddings[i].indices.tolist(),
                                values=sparse_embeddings[i].values.tolist(),
                            ),
                        },
                        payload=payload,
                    )

                points.append(point)

            # 批量写入 Qdrant
            self.qdrant_client.upsert(
                collection_name=self._get_collection_name(),
                points=points,
            )

            mode = "稀疏+稠密" if sparse_embeddings else "纯稠密"
            logger.info(f"✅ 已写入 {len(points)} 个分块到 Qdrant（{mode}）")

            # 保存分块信息到数据库
            self._save_chunks_to_db(chunks, vector_ids, document_obj)

            return vector_ids
        except Exception as e:
            logger.error(f"添加文档到向量存储失败: {e}")
            raise

    def _save_chunks_to_db(
        self,
        chunks: List[LangChainDocument],
        vector_ids: List[str],
        document_obj: Document,
    ):
        """保存分块信息到数据库"""
        chunk_objects = []
        for i, (chunk, vector_id) in enumerate(zip(chunks, vector_ids)):
            # 计算内容哈希
            content_hash = hashlib.md5(chunk.page_content.encode()).hexdigest()

            chunk_obj = DocumentChunk(
                document=document_obj,
                chunk_index=i,
                content=chunk.page_content,
                vector_id=vector_id,
                embedding_hash=content_hash,
                start_index=chunk.metadata.get("start_index"),
                end_index=chunk.metadata.get("end_index"),
                page_number=chunk.metadata.get("page"),
            )
            chunk_objects.append(chunk_obj)

        DocumentChunk.objects.bulk_create(chunk_objects)

    def similarity_search(
        self, query: str, k: int = 5, score_threshold: float = 0.1
    ) -> List[Dict[str, Any]]:
        """相似度搜索（支持稠密+稀疏混合检索）"""
        embedding_type = type(self.embeddings).__name__
        logger.info(f"🔍 开始相似度搜索 (Qdrant):")
        logger.info(f"   📝 查询: '{query}'")
        logger.info(f"   🤖 使用嵌入模型: {embedding_type}")
        logger.info(f"   🎯 返回数量: {k}, 相似度阈值: {score_threshold}")

        # 根据是否有稀疏编码器选择检索方式
        if self.sparse_encoder:
            logger.info("   🔀 使用混合检索（BM25 + 稠密向量）")
            return self._hybrid_similarity_search(query, k, score_threshold)
        else:
            logger.info("   📊 使用纯稠密向量检索")
            return self._dense_similarity_search(query, k, score_threshold)

    def _dense_similarity_search(
        self, query: str, k: int, score_threshold: float
    ) -> List[Dict[str, Any]]:
        """纯稠密向量检索"""
        try:
            dense_vector = self.embeddings.embed_query(query)
            collection_name = self._get_collection_name()

            results = self.qdrant_client.search(
                collection_name=collection_name,
                query_vector=NamedVector(
                    name=self.DENSE_VECTOR_NAME,
                    vector=dense_vector,
                ),
                limit=k,
                with_payload=True,
            )

            logger.info(f"🔍 稠密检索结果: {len(results)}")
            return self._format_search_results(results, score_threshold)

        except Exception as e:
            logger.error(f"稠密向量搜索失败: {e}")
            raise

    def _hybrid_similarity_search(
        self, query: str, k: int, score_threshold: float
    ) -> List[Dict[str, Any]]:
        """混合检索（RRF 融合稠密+稀疏 + Reranker 精排）"""
        try:
            collection_name = self._get_collection_name()
            # Reranker 需要更多候选，增加召回量
            reranker_enabled = self._get_reranker_url() is not None
            per_source_limit = max(k * 5, 20) if reranker_enabled else max(k * 3, 10)

            # 计算稠密向量
            dense_vector = self.embeddings.embed_query(query)

            # 计算稀疏向量
            sparse_query = self.sparse_encoder.encode_query(query)

            # 稠密向量检索
            dense_results = self.qdrant_client.search(
                collection_name=collection_name,
                query_vector=NamedVector(
                    name=self.DENSE_VECTOR_NAME,
                    vector=dense_vector,
                ),
                limit=per_source_limit,
                with_payload=True,
            )

            # 稀疏向量检索
            sparse_results = []
            if sparse_query:
                sparse_results = self.qdrant_client.search(
                    collection_name=collection_name,
                    query_vector=NamedSparseVector(
                        name=self.SPARSE_VECTOR_NAME,
                        vector=SparseVector(
                            indices=sparse_query.indices.tolist(),
                            values=sparse_query.values.tolist(),
                        ),
                    ),
                    limit=per_source_limit,
                    with_payload=True,
                )

            logger.info(
                f"🔍 稠密候选: {len(dense_results)}, 稀疏候选: {len(sparse_results)}"
            )

            # RRF 融合（取更多候选用于 Reranker）
            fusion_limit = k * 3 if reranker_enabled else k
            fused_results = self._rrf_fusion(
                dense_results, sparse_results, fusion_limit
            )

            # Reranker 精排（仅 Xinference 支持）
            if reranker_enabled and fused_results:
                logger.info(f"🎯 启用 Reranker 精排...")
                fused_results = self._rerank(query, fused_results, k)

            return self._format_fused_results(fused_results, score_threshold)

        except Exception as e:
            logger.error(f"混合搜索失败: {e}")
            # 降级为纯稠密检索
            logger.warning("⚠️ 降级为纯稠密检索")
            return self._dense_similarity_search(query, k, score_threshold)

    def _rrf_fusion(
        self, dense_results, sparse_results, limit: int
    ) -> List[Dict[str, Any]]:
        """RRF (Reciprocal Rank Fusion) 融合两种检索结果"""
        if not dense_results and not sparse_results:
            return []

        fused: Dict[str, Dict[str, Any]] = {}
        contributors = 0

        def accumulate(results, label: str):
            for rank, point in enumerate(results):
                point_id = str(point.id)
                if point_id not in fused:
                    fused[point_id] = {
                        "payload": point.payload or {},
                        "score": 0.0,
                        "labels": {},
                        "original_scores": {},
                    }
                incremental = 1.0 / (self.RRF_K + rank + 1)
                fused[point_id]["score"] += incremental
                fused[point_id]["labels"][label] = incremental
                fused[point_id]["original_scores"][label] = point.score

        if dense_results:
            contributors += 1
            accumulate(dense_results, "dense")
        if sparse_results:
            contributors += 1
            accumulate(sparse_results, "sparse")

        # 归一化分数到 0-1 范围
        max_possible = contributors * (1.0 / (self.RRF_K + 1))
        max_possible = max(max_possible, 1e-9)

        fused_list = []
        for point_id, data in fused.items():
            data["id"] = point_id
            data["score"] = min(data["score"] / max_possible, 1.0)
            fused_list.append(data)

        # 按融合分数降序排序
        fused_list.sort(key=lambda item: item["score"], reverse=True)
        return fused_list[:limit]

    def _format_search_results(
        self, results, score_threshold: float
    ) -> List[Dict[str, Any]]:
        """格式化稠密搜索结果"""
        formatted_results = []

        for i, point in enumerate(results):
            score = point.score
            if score < score_threshold:
                continue

            payload = point.payload or {}
            content = payload.get("page_content", "")

            result = {
                "content": content,
                "metadata": payload,
                "similarity_score": float(score),
            }
            formatted_results.append(result)


            source = payload.get("source", "未知来源")
            logger.info(
                f"   📄 结果{i + 1}: 相似度={score:.4f} ({score * 100:.1f}%), 来源={source}"
            )

        # 如果没有满足阈值的结果，返回最佳结果
        if not formatted_results and results:
            best = results[0]
            payload = best.payload or {}
            formatted_results.append(
                {
                    "content": payload.get("page_content", ""),
                    "metadata": payload,
                    "similarity_score": float(best.score),
                }
            )

        logger.info(f"📊 过滤后结果数量: {len(formatted_results)}")
        return formatted_results

    def _format_fused_results(
        self, fused_results: List[Dict], score_threshold: float
    ) -> List[Dict[str, Any]]:
        """格式化 RRF 融合结果（支持 Reranker 分数）"""
        formatted_results = []

        for i, entry in enumerate(fused_results):
            # 优先使用 rerank_score，否则使用 RRF score
            rerank_score = entry.get("rerank_score")
            score = rerank_score if rerank_score is not None else entry.get("score", 0)

            if score < score_threshold:
                continue

            payload = entry.get("payload", {})
            content = payload.get("page_content", "")

            # 添加融合来源信息
            labels = entry.get("labels", {})
            original_scores = entry.get("original_scores", {})

            result = {
                "content": content,
                "metadata": payload,
                "similarity_score": float(score),
                "fusion_detail": {
                    "sources": list(labels.keys()),
                    "dense_score": original_scores.get("dense"),
                    "sparse_score": original_scores.get("sparse"),
                    "rerank_score": rerank_score,
                },
            }
            formatted_results.append(result)

            source = payload.get("source", "未知来源")
            sources_str = "+".join(labels.keys())
            if rerank_score is not None:
                logger.info(
                    f"   📄 结果{i + 1}: Rerank分={rerank_score:.4f}, 来源={source}"
                )
            else:
                logger.info(
                    f"   📄 结果{i + 1}: 融合分={score:.4f} ({score * 100:.1f}%), 来源={source}, 检索源=[{sources_str}]"
                )

        # 如果没有满足阈值的结果，返回最佳结果
        if not formatted_results and fused_results:
            best = fused_results[0]
            payload = best.get("payload", {})
            rerank_score = best.get("rerank_score")
            score = rerank_score if rerank_score is not None else best.get("score", 0)
            formatted_results.append(
                {
                    "content": payload.get("page_content", ""),
                    "metadata": payload,
                    "similarity_score": float(score),
                }
            )

        logger.info(f"📊 过滤后结果数量: {len(formatted_results)}")
        return formatted_results

    def delete_document(self, document: Document):
        """从 Qdrant 向量存储中删除文档"""
        try:
            chunks = document.chunks.all()
            vector_ids = [chunk.vector_id for chunk in chunks if chunk.vector_id]

            if vector_ids:
                # Qdrant 删除
                collection_name = self._get_collection_name()
                self.qdrant_client.delete(
                    collection_name=collection_name, points_selector=vector_ids
                )
                logger.info(f"✅ 已从 Qdrant 删除 {len(vector_ids)} 个向量")

            chunks.delete()
        except Exception as e:
            logger.error(f"删除文档向量失败: {e}")
            raise


class KnowledgeBaseService:
    """知识库服务"""

    def __init__(self, knowledge_base: KnowledgeBase):
        self.knowledge_base = knowledge_base
        self.document_processor = DocumentProcessor()
        self.vector_manager = VectorStoreManager(knowledge_base)

    def process_document(self, document: Document) -> bool:
        """处理文档"""
        try:
            # 更新状态为处理中
            document.status = "processing"
            document.save()

            # 清理已存在的分块和向量（如果有的话）
            try:
                self.vector_manager.delete_document(document)
            except Exception as e:
                logger.warning(f"删除旧向量时出错（可能是首次处理）: {e}")

            # 再从数据库删除分块记录
            document.chunks.all().delete()

            # 加载文档
            langchain_docs = self.document_processor.load_document(document)

            # 计算文档统计信息
            total_content = "\n".join([doc.page_content for doc in langchain_docs])
            document.word_count = len(total_content.split())
            document.page_count = len(langchain_docs)

            # 向量化并存储
            vector_ids = self.vector_manager.add_documents(langchain_docs, document)

            # 更新状态为完成
            document.status = "completed"
            document.processed_at = timezone.now()
            document.error_message = None
            document.save()

            logger.info(f"文档处理成功: {document.id}, 生成 {len(vector_ids)} 个分块")
            return True

        except Exception as e:
            # 更新状态为失败
            document.status = "failed"
            document.error_message = str(e)
            document.save()

            logger.error(f"文档处理失败: {document.id}, 错误: {e}")
            return False

    def query(
        self,
        query_text: str,
        top_k: int = 5,
        similarity_threshold: float = 0.5,
        user=None,
    ) -> Dict[str, Any]:
        """查询知识库"""
        start_time = time.time()

        try:
            # 记录查询开始信息
            embedding_type = type(self.vector_manager.embeddings).__name__
            logger.info(f"🚀 知识库查询开始:")
            logger.info(f"   📚 知识库: {self.knowledge_base.name}")
            logger.info(f"   👤 用户: {user.username if user else '匿名'}")
            logger.info(f"   🤖 嵌入模型: {embedding_type}")
            logger.info(f"   💾 向量存储: Qdrant")

            # 执行检索
            retrieval_start = time.time()
            search_results = self.vector_manager.similarity_search(
                query_text, k=top_k, score_threshold=similarity_threshold
            )
            retrieval_time = time.time() - retrieval_start

            # 生成回答（这里可以集成LLM）
            generation_start = time.time()
            answer = self._generate_answer(query_text, search_results)
            generation_time = time.time() - generation_start

            total_time = time.time() - start_time

            # 记录查询日志
            self._log_query(
                query_text,
                answer,
                search_results,
                retrieval_time,
                generation_time,
                total_time,
                user,
            )

            # 记录查询完成信息
            logger.info(f"✅ 知识库查询完成:")
            logger.info(f"   ⏱️  检索耗时: {retrieval_time:.3f}s")
            logger.info(f"   🤖 生成耗时: {generation_time:.3f}s")
            logger.info(f"   🕐 总耗时: {total_time:.3f}s")
            logger.info(f"   📊 返回结果数: {len(search_results)}")

            return {
                "query": query_text,
                "answer": answer,
                "sources": search_results,
                "retrieval_time": retrieval_time,
                "generation_time": generation_time,
                "total_time": total_time,
            }

        except Exception as e:
            logger.error(f"知识库查询失败: {e}")
            raise

    def _generate_answer(self, query: str, sources: List[Dict[str, Any]]) -> str:
        """生成回答（简单版本，后续可集成LLM）"""
        if not sources:
            return "抱歉，没有找到相关信息。"

        # 简单的基于检索结果的回答生成
        context = "\n\n".join([source["content"] for source in sources[:3]])
        return f"基于查询「{query}」检索到的相关内容：\n\n{context}"

    def _log_query(
        self,
        query: str,
        answer: str,
        sources: List[Dict[str, Any]],
        retrieval_time: float,
        generation_time: float,
        total_time: float,
        user,
    ):
        """记录查询日志"""
        try:
            QueryLog.objects.create(
                knowledge_base=self.knowledge_base,
                user=user,
                query=query,
                response=answer,
                retrieved_chunks=[
                    {
                        "content": source["content"][:200] + "..."
                        if len(source["content"]) > 200
                        else source["content"],
                        "metadata": source["metadata"],
                        "score": source["similarity_score"],
                    }
                    for source in sources
                ],
                similarity_scores=[source["similarity_score"] for source in sources],
                retrieval_time=retrieval_time,
                generation_time=generation_time,
                total_time=total_time,
            )
        except Exception as e:
            logger.error(f"记录查询日志失败: {e}")

    def delete_document(self, document: Document):
        """删除文档"""
        try:
            # 从向量存储中删除
            self.vector_manager.delete_document(document)

            # 删除文件
            if document.file:
                if os.path.exists(document.file.path):
                    os.remove(document.file.path)

            # 删除数据库记录
            document.delete()

            # 清理向量存储缓存（因为内容已变化）
            VectorStoreManager.clear_cache(self.knowledge_base.id)

            logger.info(f"文档删除成功: {document.id}")

        except Exception as e:
            logger.error(f"删除文档失败: {e}")
            raise
